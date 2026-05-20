"""
Word 文档解析
"""
from typing import Optional
from docx import Document
from core.parser.base import BaseParser


class DocxParser(BaseParser):
    """Word .docx 文件解析器"""

    @property
    def file_type(self) -> str:
        return ".docx"

    def parse(self, file_path: str) -> str:
        """解析文件"""
        doc = Document(file_path)
        return self._extract_text(doc)

    def parse_content(self, content: bytes) -> str:
        """解析文件内容"""
        import io
        doc = Document(io.BytesIO(content))
        return self._extract_text(doc)

    def _extract_text(self, doc: Document) -> str:
        """从Document对象提取文本"""
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n".join(paragraphs)


# 全局实例
_docx_parser: Optional[DocxParser] = None


def get_docx_parser() -> DocxParser:
    """获取解析器"""
    global _docx_parser
    if _docx_parser is None:
        _docx_parser = DocxParser()
    return _docx_parser
